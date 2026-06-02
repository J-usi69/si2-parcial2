"""Exportacion de un reporte (titulo, columnas, filas) a Excel, Word y PDF."""
import io


def to_xlsx(title: str, columns: list[str], rows: list[list]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = "Reporte"

    ws.append([title])
    ws["A1"].font = Font(bold=True, size=14)
    ws.append([])

    header_fill = PatternFill(start_color="1E88E5", end_color="1E88E5", fill_type="solid")
    ws.append(columns)
    header_row = ws.max_row
    for col_idx in range(1, len(columns) + 1):
        cell = ws.cell(row=header_row, column=col_idx)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill

    for row in rows:
        ws.append(["" if v is None else v for v in row])

    # Ancho de columnas aproximado.
    for col_idx, name in enumerate(columns, start=1):
        width = max(len(str(name)), 12)
        ws.column_dimensions[ws.cell(row=header_row, column=col_idx).column_letter].width = min(width + 2, 40)

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def to_docx(title: str, columns: list[str], rows: list[list]) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.size = Pt(18)

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    for i, name in enumerate(columns):
        cell = table.rows[0].cells[i]
        cell.text = str(name)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def to_pdf(title: str, columns: list[str], rows: list[list]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), title=title)
    styles = getSampleStyleSheet()
    elements = [Paragraph(title, styles["Title"]), Spacer(1, 12)]

    # Limitar columnas/filas para que quepan en PDF.
    max_cols = 8
    cols = columns[:max_cols]
    truncated_note = len(columns) > max_cols

    def _fmt(v):
        s = "" if v is None else str(v)
        return s if len(s) <= 40 else s[:37] + "..."

    data = [cols] + [[_fmt(v) for v in row[:max_cols]] for row in rows[:200]]
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E88E5")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F6FC")]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elements.append(table)
    if truncated_note:
        elements.append(Spacer(1, 8))
        elements.append(Paragraph(f"(Mostrando {len(cols)} de {len(columns)} columnas)", styles["Italic"]))

    doc.build(elements)
    return buffer.getvalue()


EXPORTERS = {
    "xlsx": (to_xlsx, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
    "docx": (to_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "pdf": (to_pdf, "application/pdf"),
}
